import streamlit as st
import argparse
import random
import os
import json
import subprocess
from re import sub
import re
from typing import AsyncContextManager
from io import StringIO
from html.parser import HTMLParser
import pprint 
import numpy as np
import stripe
import docx
import openai
import pandas as pd
import uuid
stripe_keys = {
    "secret_key": os.environ["STRIPE_SECRET_KEY"],
    "publishable_key": os.environ["STRIPE_PUBLISHABLE_KEY"],
    "price_id": os.environ["STRIPE_PRICE_ID"],
    "endpoint_secret": os.environ["STRIPE_ENDPOINT_SECRET"],
}
stripe.api_key = stripe_keys["secret_key"]
from app.utilities.gpt3_reduced_fx import gpt3complete, presets_parser
def create_uuid():
    return str(uuid.uuid4())
class MLStripper(HTMLParser):
    def __init__(self):
        super().__init__()
        self.reset()
        self.strict = False
        self.convert_charrefs= True
        self.text = StringIO()
    def handle_data(self, d):
        self.text.write(d)
    def get_data(self):
        return self.text.getvalue()
   
    
    
def strip_tags(html):
    s = MLStripper()
    s.feed(html)
    return s.get_data()
def post_process_text(text):
    #all_patterns = [r'<br\s*?>', r'<br>', r'<li>', r'\n\s*\n', r'^-\s+', r'^-', r'\d+[)]'
    #combined_patterns =  = r'|'.join(map(r'(?:{})'.format, all_pats))
    text = text.replace('<br\s*?>', '')
    text = text.replace('<br>', '\n')
    text = text.replace('<li>', '\n')
    text = re.sub(r'\d+[)]', "", text)
    text = text.replace('\n-', '\n')
    text = text.replace('\n• ','\n')
    text = text.replace('\n ', '\n')
    text = text.replace('\n\n', '\n')
    text = text.replace('\"\"', '\"')
    text = re.sub('\d+[.]\s+', '', text).rstrip()
    text = re.sub('\d+[.\n]\s+', '', text).rstrip()
   #print('post processed text is', '\n' + text)
    return text
def sort_and_uniq_pool(pool):
    sortedpool = sorted((pool).split('\n'))
    pool = '\n'.join([i for n, i in enumerate(sortedpool) if i not in sortedpool[:n]])
   #print("sorted pool", pool)
    return pool
def vertical_chooser(vertical):
    
    if vertical == '':
        pass
    elif vertical == 'random':
            valid_industries = ["eCommerce", "Retail", "Consumer Goods", "B2C SaaS", "Entertainment"]
            vertical = random.choice(valid_industries)
    else:
         vertical = vertical 
   #print('returning vertical', vertical)
    return vertical
def choose_user_type():
    valid_user_types = ['Free', 'Professional', 'Enterprise']
    user_type = random.choice(valid_user_types)
    return user_type 
def create_record_topics(datadir, record_topic_preset, n=20, insert_before_post_user_input="", override_post="", vertical="", user_type=""):
    # if request is coming from api
    #datadir = 'app/data/apitest'
    remove_html = True
    rawrecordtopicsfilepath = datadir + '/' + 'raw_record_topics.txt'
    try:
        os.remove(rawrecordtopicsfilepath)
    except:
       print("Error while deleting file ", rawrecordtopicsfilepath)
        
    # read preset to get default values
    
    presetsdf, preset_name, preset_description, preset_instructions, preset_additional_notes, preset_placeholder, pre_user_input, post_user_input, prompt, engine, finetune_model, temperature, max_tokens, top_p, fp, pp, stop_sequence, echo_on, preset_pagetype, preset_db, user, organization = presets_parser(record_topic_preset)
    #print(presetsdf)
    if insert_before_post_user_input:
        additional_post_user_input = '\n'
    lines = 0
    record_topic = []
   
    max_tokens = 180
    stop_sequence_no = int(13 - lines)
    stop_sequence = str(stop_sequence_no) + '.' +' '
    n = int(n)
    #print(n)
    
    # read prompt inputs from function parameters
    
        
    if vertical == 'Random':
        vertical_prompt = vertical_chooser(vertical)
    
    elif vertical:
        vertical_prompt = "The company is in the " + vertical + " industry."
        
    else:
        vertical_prompt = ""
    if override_post:
       post_user_input = override_post
    
        
    #st.write('post user input in CSD is ', post_user_input)
    prompt = pre_user_input + prompt + insert_before_post_user_input + post_user_input + vertical_prompt
    st.write('Final prompt submitted to language generation model was: ')
    st.write(prompt)
    report = "record topic prompt: " + prompt
    recordtopicpool = ""
    recordtopiclist = []
    listoftopics = []
    while lines < n:
    
        response = openai.Completion.create(
        engine=engine,
        prompt=prompt,
        temperature=temperature,
        max_tokens=max_tokens,
        top_p=top_p,
        frequency_penalty=fp,
        presence_penalty=pp,
        logprobs=0,
        echo=echo_on,
        stop = stop_sequence,
        user="testing", 
        organization = organization
    )   
        #print(response)
        response_text = response['choices'][0]['text']
        #print('response with lines', lines, response_text)
  
        # writing raw response  for debugging purposes
        rawrecordtopicsfilepath = datadir + '/' + 'raw_record_topics.txt'
        with open(rawrecordtopicsfilepath, 'a+') as f:
            f.write(response_text + '\n')
        # now beginning to process raw response
        if remove_html:
            response_text = strip_tags(response_text)
        else:
            pass # allows html tags in response
        s = post_process_text(response_text)
        recordtopicpool = recordtopicpool + s
        lines = recordtopicpool.count("\n") + 1
        recordtopiclist = sorted(recordtopicpool.split('\n'))
        recordtopicpool = sort_and_uniq_pool(recordtopicpool)
        recordtopicpath = datadir + '/' 'recordtopics.txt'
        with open(recordtopicpath, 'w') as f:
                f.write(recordtopicpool)
       #print('wrote record topics to ', recordtopicpath)
    
    return vertical, user_type, recordtopiclist, report
def create_status():
    valid_statuses = ["New", "Open", "Pending", "Solved"]
    status = random.choice(valid_statuses)
    # assign status based on content
    return status
def create_one_record(datadir, vertical, user_type, record_topic, status, customer_name, agent_name, record_max_length, record_writer_preset):
   #print(record_writer_preset)
    presetsdf, preset_name, preset_description, preset_instructions, preset_additional_notes, preset_placeholder, pre_user_input, post_user_input, prompt, engine, finetune_model, temperature, max_tokens, top_p, fp, pp, stop_sequence, echo_on, preset_pagetype, preset_db, user, organization = presets_parser(record_writer_preset)
    if finetune_model:
        engine = finetune_model
    vertical = vertical_chooser(vertical)
    
    if vertical:
        vertical_prompt = "The company is in the " + vertical + " industry."
    else:
        vertical_prompt = ""
    prompt=pre_user_input + prompt + record_topic + '\n' + post_user_input + vertical_prompt
   #print('create record prompt',prompt)
    response = openai.Completion.create(
        engine=engine,
        prompt=prompt,
        temperature=temperature,
        max_tokens=max_tokens,
        top_p=top_p,
        frequency_penalty=fp,
        presence_penalty=pp,
        logprobs=0,
        echo=echo_on,
        stop = stop_sequence,
        user="testing",
        organization = organization
    )
    result = "Record Topic: " + record_topic + '\n' + "[Status] " + status + '\n'
    response_text = response['choices'][0]['text']
   #print('create record response', response_text)
    allrawrecordsfile = datadir + '/' + 'allrawrecords.txt'
    try:
        os.remove(allrawrecordsfile)
    except:
       print("Error while deleting file ", allrawrecordsfile)
    with open(allrawrecordsfile, 'a+')  as f:
            f.write(response_text)
    if remove_html:
        response_text = strip_tags(response_text)
    response_text = post_process_text(response_text).lstrip()
    record =  result + response_text
    return record, user_type 
def classify_record(record, user_type):
    # addingmore fine-grained classification levels
    # atlassian_sev_levels = {[{"name" : "Severity 1", "description": "A critical incident with very high impact"},{"name": "Severity 2", "description" : "A major incident with significant impact"}, {"name": "Severity 3", "description" : "A minor incident with low impact"}]}
    lnh = ["Low", "Normal", "High"] 
    priority_config = lnh
    valid_priorities = priority_config
    presetsdf, preset_name, preset_description, preset_instructions, preset_additional_notes, preset_placeholder, pre_user_input, post_user_input, prompt, engine, finetune_model, temperature, max_tokens, top_p, fp, pp, stop_sequence, echo_on, preset_pagetype, preset_db, user, organization = presets_parser(classifier_prompt_preset)
    if finetune_model:
        engine = finetune_model
   #print ("classifying one record")
    
    prompt=record[0]
    
    prompt=pre_user_input + prompt + post_user_input
   #print('classification prompt is', prompt)
    response = gpt3complete(
    engine="davinci-instruct-beta",
    prompt=prompt,
    temperature=temperature,
    max_tokens=max_tokens,
    top_p=top_p,
    frequency_penalty=fp,
    presence_penalty=pp,
    logprobs=0,
    user="testing",
    organization = organization
    )
    response_text = response['choices'][0]['text']
   #print('raw priority', response_text)
    priority = post_process_text(response_text).rstrip()
   #print('priority pp', priority)
    prioritized = record[0]  + '\n' + "[Priority]" + priority + '\n' + '[$AccType]' + user_type + '\n####' 
    return prioritized
def create_article_titles(datadir, vertical,  n, article_title_preset, article_title_prompt):
    lines = 0
    kbtitle = []
    kbpool = ""
    max_tokens = 100
    if article_title_prompt:
        override_prompt = article_title_prompt

    presetsdf, preset_name, preset_description, preset_instructions, preset_additional_notes, preset_placeholder, pre_user_input, post_user_input, prompt, engine, finetune_model, temperature, max_tokens, top_p, fp, pp, stop_sequence, echo_on, preset_pagetype, preset_db, user, organization = presets_parser(article_title_preset)

    #print('preset_name', preset_name, 'preset_description', preset_description, 'preset_instructions', preset_instructions, 'preset_placeholder', preset_placeholder, 'pre_user_input', pre_user_input, 'post_user_input', post_user_input, 'prompt', prompt, 'engine', engine, 'finetune_model', finetune_model, 'temperature', temperature, 'max_tokens', max_tokens, 'top_p', top_p, 'fp', fp, 'pp', pp, 'stop_sequence', stop_sequence, 'echo_on', echo_on, 'preset_db', preset_db, 'user', user, 'organization', organization)

    #vertical = vertical_chooser(vertical)
    
    if vertical:
        vertical_prompt = vertical
    else:
        vertical_prompt = ""
    if override_prompt:
        prompt = prompt + '\n' + override_prompt
    #print('pre_user_input', pre_user_input, '\nprompt', prompt, '\npost_user_input', post_user_input, '\nvertical_prompt', vertical_prompt)
    prompt = vertical_prompt + pre_user_input + prompt + post_user_input 
    print('*****')
    print('kb title prompt is:\n', prompt)
    print('*****')
    n = int(n)
    while lines < n:
    
        response = openai.Completion.create(
        engine=engine,
        prompt=prompt,
        temperature=temperature,
        max_tokens=max_tokens,
        top_p=top_p,
        frequency_penalty=fp,
        presence_penalty=pp,
        logprobs=0,
        echo=echo_on,
        stop = stop_sequence,
        user="testing",
        organization = organization
    )
        #print(response)
        response_text = response['choices'][0]['text']
        #print("generated response:", response_text)
        # do bespoke postprocessing
        with open(datadir + '/' + 'kbrawtitles.txt', 'a')  as f:
            f.write("###\n" + response_text + "\n")
        s = post_process_text(response_text)
        kbpool = kbpool + s
        lines = kbpool.count("\n") + 1
        kbpool = sort_and_uniq_pool(kbpool)
        kbtitlepath = datadir + '/' 'kbtitles.txt'
        with open(kbtitlepath, 'w') as f:
                f.write(kbpool)
       #print('wrote article titles to ', kbtitlepath)
    return vertical, kbpool, kbtitlepath
def create_kb_article(vertical, kbtitle, article_writer_preset, article_length, article_length_random, article_writer_prompt):
    if article_length_random:
        length_choices = [100, 200, 300, 400, 600, 1000]
        article_length = random.choice(length_choices)
    else:
        article_length = article_length
    if article_writer_prompt:
        override_prompt = article_writer_prompt
   #print("creating  article, length " + str(article_length) + " tokens")
   #print('title of article is', kbtitle)
   #print('article preset is', article_writer_preset)
    presetsdf, preset_name, preset_description, preset_instructions, preset_additional_notes, preset_placeholder, pre_user_input, post_user_input, prompt, engine, finetune_model, temperature, max_tokens, top_p, fp, pp, stop_sequence, echo_on, preset_pagetype, preset_db, user, organization, vertical = presets_parser(article_writer_preset)
    vertical = vertical_chooser(vertical)
    
    if vertical:
        vertical_prompt = " The company is in the " + vertical + " industry."
    else:
        vertical_prompt = ""
    if override_prompt:
        prompt = override_prompt
    
    final_prompt = pre_user_input + prompt + post_user_input
   #print(pre_user_input + prompt + post_user_input + vertical_prompt)
   #print('engine is', engine)
    post_user_input =  "\nTitle:\n" + kbtitle + '\n' + post_user_input + vertical_prompt
    response = openai.Completion.create(
    engine=engine,
    prompt=final_prompt,
    temperature=temperature,
    max_tokens=article_length,
    top_p=top_p,
    stop=stop_sequence,
    frequency_penalty=fp,
    presence_penalty=pp,
    user="testing",
    organization=organization,
    echo=True,
    )
    #print(response)
    response_text = response['choices'][0]['text']
    #print(response_text)
    article = "Title: " + kbtitle + '\n' + response_text 
    return article
def batch_run(datadir, batchfile):
    import json
    with open(batchfile, 'r') as f:
        batchdata = json.load(f)
    
    vertical_tuples = batchdata['vertical_tuples']
    for vertical in vertical_tuples:
       #print(vertical)
        pass
    return
def main(article_length=300, article_length_random=False, article_number=3, article_title_prompt='intelligent tanks', article_title_input_file="", article_title_preset='SimpleXmasStoryIdeas', article_titles_only=True, article_writer_preset='kbTopicA', article_writer_prompt='imaginative', batchfile=False, classifier_prompt_preset='recordclassifierA', datadir='/tmp/longform', finetune_model='ada', generate_articles=True, generate_records=False, insert_before_post_user_input="", number=5, record_prompt_list=['crmA', 'crmE'], record_topic_preset='GenericGenerateTicketTopics', record_topics_only=False, record_writer_preset='GenericRecordWriter', remove_html=True, user_type="system user", vertical=""):
   
    if not os.path.exists(datadir):
        os.makedirs(datadir)
       #print('created datadir', datadir)
    else:
       print("directory already exists, using existing directory")
    if batchfile:
       #print('generate records', generate_records, 'generate articles', generate_articles, 'run batch job', batchfile)
        batch_run(datadir, batchfile)
    else:
        pass
    report = ""
    if not generate_records:
       pass#print("not generating records")
    else:
       #print('generating records')
        if user_type == "random":
            user_type = choose_user_type()
        else:
            pass
    
        #record_prompt_preset = np.random.choice(["crmG", "crmG"], p=[0.3, 0.7])
       #print('chose preset', record_topic_preset)
       #print('datadir is', datadir)
        recordtopics = create_record_topics(datadir, vertical, user_type, record_topic_preset, number, insert_before_post_user_input)
        pool = recordtopics[2]
        report = report + recordtopics[3]
        if record_topics_only:
           pass#print ('only generated record topics, did not create records')
        else:
            records=[]
           #print("now entering main loop to create records")
            for record_topic in pool.splitlines():
               #print("subject line is", record_topic)
                status = create_status()
                record = create_one_record(datadir, vertical, user_type, record_topic, status, "john doe", "Mari V.",  200, record_writer_preset)
                prioritized = classify_record(record, user_type)
                records.append(prioritized)
                #print('returned records are', records)
                recordspath = datadir + '/' + 'records.txt'
            with open(recordspath, 'w') as f:
                for item in records:
                    f.write("%s\n" % item)
                f.close()   
    if not generate_articles:
       print("not generating articles")
    else:
        kbpath = datadir + '/' + 'GeneratedKBarticles.txt'
        kbdocx = datadir + '/' + 'GeneratedKBarticles.docx'
       #print('generating article titles for ' + vertical + ' vertical using  article_title_preset')
        if article_title_input_file:
            with open(article_title_input_file, 'r') as f:
                kbpool = f.read()
            kbtitlepath=article_title_input_file
            report = report + 'article titles read from file ' + kbtitlepath + '\n'
        else: 
            kbcreate = create_article_titles(datadir, vertical, article_number, article_title_preset, article_title_prompt)
            kbpool = kbcreate[1]
            kbtitlepath = kbcreate[2]
           #print('creating article titles de novo')
            report = report + '\n' + 'wrote raw article titles to ' +  datadir + '/' + 'kbrawtitles.txt'
            report = report + '\n' + ' wrote sorted & deduped article titles to '  + kbtitlepath
        
        ##print('kbpool is', kbpool)
       #print('article titles only =', article_titles_only)
        if article_titles_only:
           #print('titles only', kbpool)
            pass
        else:
           #print('writing articles')
        
            # first clear up old files
            try:
                os.remove(kbpath)
            except:
               print("Error while deleting file ", kbpath)
            # loop through kbpool to build articles
            doc = docx.Document()
            nlines = 0
            for kbtitle in kbpool.splitlines():
                article = create_kb_article(vertical, kbtitle, article_writer_preset, article_length, article_length_random, article_writer_prompt) + "\n####\n"
                short_article_title=False
                if short_article_title:
                    kbtitle = kbtitle.split(':')[0]
                else:
                    pass
                paragraph = doc.add_paragraph(kbtitle, style='Heading 2')
               #print("article text: ", article)
                with open(kbpath, 'a+') as f:
                        f.write(article)
                        f.close()
                paragraph = doc.add_paragraph(article,style='Body Text')
                nlines = len(article.splitlines()) + nlines
      
            doc.save(kbdocx)
    # wrap up report
    report = report + record_topic_preset
    report = report + '\n' + record_writer_preset
    report = report + '\n' + article_title_preset
    report = report + '\n' + article_writer_preset
    if generate_records:
        report = report + 'wrote record topics to ' + datadir + '/' + 'pp_record_topics.txt'
        if not record_topics_only:
            report = report + 'wrote records to ' + recordspath
    else:
        report = report + 'did not generate records'
    if generate_articles:
        if not article_titles_only:
            report = report + '\n' + 'wrote word doc to '+ kbdocx
            report = report + '\n' + 'wrote articles to '  +  kbpath
            report = report + '\n' + 'there were ' +  str(nlines) + ' lines in the articles'       
    else:
        report = report + '\n' + 'did not generate articles'
   #print(report)
    filesuffix = vertical.replace(" ", "_")
    filename = datadir + '/' + 'README_' + filesuffix + '.md'
    with open(filename, 'w') as f:
        f.write(report)
        f.close()
   #print('README is at ', filename)
uuid = create_uuid()
print('uuid is', uuid)